from PaPDF import PaPDF
import sys
import os
from docx import Document
import time
from system_notes import Venezuela, Colombia
from time import sleep
from PyQt5 import uic, QtWidgets, QtGui, QtCore
from PyQt5.QtWidgets import *
from PyQt5.QtGui import *
from PyQt5.QtCore import *

#Funciones
def crear_carpeta(nombre):

    #Directorio
    parent_dir = r"C:\Users\Juan\Documents"

    #Aqui obtenemos la ruta
    path = os.path.join(parent_dir, nombre)

    #Verificar si el directorio existe
    if os.path.isdir(path) == False:
        #Creamos la carpeta
        os.mkdir(path)

    else:
        pass

##########

#Esta clase nos permite guardar los valores de las configuraciones
setting = QSettings("Notes", "Config")

qtCreatorFile = os.path.abspath("ui/notes.ui")

Ui_MainWindow, QtBaseClass = uic.loadUiType(qtCreatorFile)



class VentanaPrincipal(QtWidgets.QMainWindow, Ui_MainWindow):
    def __init__(self):
        QtWidgets.QMainWindow.__init__(self)
        Ui_MainWindow.__init__(self)
        self.setupUi(self)

        #Se mueve la ventana a esas coordenadas
        self.move(150, 0)

        #Se establece un ancho y alto de la ventana definitivo
        self.setFixedSize(700,500)

        #Conectamos el evento de los botones a las dos funciones de abajo
        self.inicio.clicked.connect(self.notes)
        self.opciones.clicked.connect(self.show_opciones)

        #Creamos las clases de las dos ventanas
        #Estas se abriran al presionar un boton
        self.notes = VentanaNotes()
        self.opciones_menu = Ventana_Opciones()

        #Sombre de el boton de inicio
        self.shadow_inicio = QtWidgets.QGraphicsDropShadowEffect()
        self.shadow_inicio.setBlurRadius(15)

        #Sombre de el boton de opciones
        self.shadow_opciones = QtWidgets.QGraphicsDropShadowEffect()
        self.shadow_opciones.setBlurRadius(15)

        #Agregamos las dos sombras a sus botones 
        self.inicio.setGraphicsEffect(self.shadow_inicio)
        self.opciones.setGraphicsEffect(self.shadow_opciones)



    def notes(self):

        #Primero ejecutamos las configuraciones
        #Antes de abrir la ventana
        self.notes.config_exec()

        #Se abre la ventana
        self.notes.exec_()


    def show_opciones(self):

        #Se abre la ventana
        self.opciones_menu.exec_()



class VentanaNotes(QDialog):
    def __init__(self):
        QDialog.__init__(self)
        uic.loadUi(os.path.abspath("ui/notes_2.ui"), self)
        
        self.setFixedSize(651, 500)
        self.move(175, 10)


        self.shadow_prueba = QtWidgets.QGraphicsDropShadowEffect()
        self.shadow_prueba.setBlurRadius(110)
        self.frame.setGraphicsEffect(self.shadow_prueba)
        
        self.anadir.clicked.connect(self.texto)
        self.button_promedio.clicked.connect(self.promedio)

        #Creamos la barra deslizante
        scroll = QScrollBar()

        #Y se añade a la listbox
        self.lista.setVerticalScrollBar(scroll)

    def config_exec(self):

        try:
            #Obtenemos los valores guardados
            self.exportacion = setting.value("exportacion")
            self.sistema_notas = setting.value("sistema_notas")
            self.numero_notas = setting.value("numero_notas")

            #Ejecutara el sistema de notas que se encuentre en la configuracion
            if self.sistema_notas == "Venezuela-Liceo":

                self.sistema_notas = Venezuela()
                self.nota.setText("0")

            else:

                self.sistema_notas = Colombia()
                self.nota.setText("0.0")

            
        except:
            #Se dara valores por defectos
            #En caso de que no se hayan especificados
            self.sistema_notas = Colombia()
            self.numero_notas = 5
            self.exportacion = True


         
    def promedio(self):
        global lista_notas, promedio_text

        #Verificamos si el numero de notas es menor al limite
        if (self.lista.count() < self.numero_notas):
            QtWidgets.QMessageBox.warning(self, "Error_numero_notas", "Se necesitan el numero de notas exactas al limite", QtWidgets.QMessageBox.Ok)

        else:
            #Se obtendran todas las notas de la listbox
            lista_notas = [str(self.lista.item(i).text()[9:12]) for i in range(self.lista.count())]

            #Y se obtiene el promedio de notas
            promedio_text = self.sistema_notas.promedio(self.numero_notas, lista_notas)

            #Y se muestra el promedio
            self.nota.setText(str(promedio_text))


            #Verificamos si esta activado la exportacion
            if self.exportacion == "Si":

                #En ese caso se mostrara la ventana en un par de segundos
                time.sleep(7.8)
                self.ventana_export = VentanaExportacion()
                self.ventana_export.exec_()

            elif self.exportacion == "No":
                pass

                
                
    def texto(self):

        #Verificamos si ya estamos al limite de las notas permitidas
        if (self.lista.count() == self.numero_notas):
            QtWidgets.QMessageBox.warning(self, "Error_numero_notas", "No se puede sobrepasar el limite de notas", QtWidgets.QMessageBox.Ok)

            #Y borramos la nota que se iba agregar de la entrada
            self.entrada.setText("")

        else:

            #Convertimos en string la nota que se encuentra en la entrada
            self.nota_texto = str(self.entrada.text())

            try:

                #Dependiendo de la nota se pondra un icono diferente
                #Indicando si es una nota buena, mala o intermedia
                self.icon_use = QIcon(self.sistema_notas.texto(self.nota_texto))

                #Se usa un contador para poner un numero a cada nota
                count_lista = str(self.lista.count() + 1)

                #Se crea un objeto item listbox
                elemento = QListWidgetItem(f"Nota {count_lista} - {self.nota_texto}", self.lista)

                #Se alinea el item al centro de la listbox
                elemento.setTextAlignment(Qt.AlignCenter)

                #Y se modifica el tamaño del icono para que sea mas grande
                self.lista.setIconSize(QSize(40, 40))

                #Se agrega el icono al item listbox
                elemento.setIcon(self.icon_use)
                    
                #Se agrega el item a la listbox
                self.lista.addItem(elemento)

                #Y la entrada queda vacia
                self.entrada.setText("") 
                
            except ValueError:
                #Si no es posible convertir el contenido de la entrada en un float sera invalida
                QtWidgets.QMessageBox.warning(self, "Error_contenido", "La informacion ingresada no es valida", QtWidgets.QMessageBox.Ok)

        
class VentanaExportacion(QDialog):
    def __init__(self):
        QDialog.__init__(self)
        uic.loadUi(os.path.abspath("ui/exportacion.ui"), self)

        self.move(320,95)

        self.word.clicked.connect(self.word_execute)
        self.pdf.clicked.connect(self.pdf_execute)

    def word_execute(self):
        #Aqui se crea el documento word

        self.doc = Document()

        #Se agregan dos titulos
        self.doc.add_heading("Promedio de Notas", 0)
        self.doc.add_heading(f"Promedio: {promedio_text}", level=1)

        #Se escriben todas las notas
        for nota in lista_notas:
            self.doc.add_paragraph(f"{nota}", style="List Bullet")

        crear_carpeta("Promedios")

        #Se cierra la pagina de word
        self.doc.add_page_break()

        self.doc.save(f"C:/Users/Juan/Documents/Promedios/Promedio {self.count_file}.docx")
        self.close()


    def pdf_execute(self):
        #Aqui se crea el documento PDF


        with PaPDF(f"C:/Users/Juan/Documents/Promedios/Promedio {self.count_file}.pdf") as pdf:

            #Titulo
            pdf.setFontSize(25)
            pdf.addText(54, 290, "Promedio Notas")

            #Promedio Total
            pdf.setFontSize(20)
            pdf.addText(60, 250, f"Promedio: {promedio_text}")

            #Lista de Notas
            count_pos = 235
            for nota in lista_notas:
                count_pos-=5
                pdf.setFontSize(17)
                pdf.addText(75, count_pos, nota)

            pdf.addPage()
            self.close()




class Pantalla_Carga(QtWidgets.QMainWindow):
    def __init__(self):
        QtWidgets.QMainWindow.__init__(self)
        uic.loadUi(os.path.abspath("ui/pantalla_carga.ui") , self)

        self.setWindowTitle("Pantalla de carga")

        #Esta clase sirve para crear una sombra
        shadow = QGraphicsDropShadowEffect()
        shadow.setBlurRadius(25)

        #Agregamos la sombre al icono
        self.icono.setGraphicsEffect(shadow)

        #Agregamos un icono al ToolButton y agregramos la ruta del icono
        self.icono.setIcon(QIcon("imagenes/-pen_96740.ico"))

        #Asi borramos tanto el marco de ventana, como el fondo
        self.setWindowFlag(Qt.FramelessWindowHint)
        self.setAttribute(Qt.WA_TranslucentBackground)

        #Configuramos el tiempo de la barra de carga
        self.reloj = QTimer()
        self.reloj.timeout.connect(self.barra_carga)
        self.reloj.start(45)

        self.show()

    def barra_carga(self):
        sleep_time = 0.09
        for i in range(1, 101):
            sleep(0.09)
            self.progressBar.setValue(i)

        self.reloj.stop()
        self.close()
        self.ventana_main = VentanaPrincipal()
        self.ventana_main.show()


class Ventana_Opciones(QDialog):
    def __init__(self):
        QDialog.__init__(self)

        uic.loadUi(os.path.abspath("ui/opciones.ui"), self)

        self.setWindowTitle("Menu Opciones")

        #Creamos la sombra
        shadow_icono = QGraphicsDropShadowEffect()
        shadow_icono.setBlurRadius(25)

        #Se añade las sombras al icono 
        self.icono.setGraphicsEffect(shadow_icono)

        #Aqui añadimos los iconos de el pais correspondiente al sistema de notas
        self.sistema_notas.addItem(QIcon("imagenes/pais_sistema/venezuela.ico"), "Venezuela-Liceo")
        self.sistema_notas.addItem(QIcon("imagenes/pais_sistema/colombia.ico"), "Colombia")

        #Añadimos el numero de notas permitidos
        self.numero_notas.addItem("5")
        self.numero_notas.addItem("10")
        self.numero_notas.addItem("15")
        self.numero_notas.addItem("20")

        #Aqui las dos opciones de exportacion
        self.exportacion.addItem("Si")
        self.exportacion.addItem("No")

        #Aqui añadimos los eventos de los botones "cancelar" y "aplicar"
        self.aplicar.clicked.connect(self.configuracion)
        self.cancelar.clicked.connect(self.salir)

    def salir(self):
        self.close()

    def configuracion(self):

        #Obtenemos la informacion elegida del combox
        self.content_sys = self.sistema_notas.currentText()
        self.content_export = self.exportacion.currentText()
        self.content_numero = int(self.numero_notas.currentText())

        #Aqui guardamos las configuraciones
        setting.setValue("sistema_notas", self.content_sys)
        setting.setValue("exportacion", self.content_export)
        setting.setValue("numero_notas", self.content_numero)
        self.close()




        




        

      









if __name__ == "__main__":
    app =  QtWidgets.QApplication(sys.argv)
    window = Pantalla_Carga()
    sys.exit(app.exec_())
